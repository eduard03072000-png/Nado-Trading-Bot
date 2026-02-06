"""
Генератор Word отчетов с аналитикой торговли
Автоматическое создание и обновление отчетов
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """
    Класс для генерации аналитических отчетов в Word
    """
    
    def __init__(self, reports_dir: str = "data/reports"):
        self.reports_dir = Path(reports_dir)
        self.reports_dir.mkdir(parents=True, exist_ok=True)
    
    def create_daily_report(
        self,
        database,
        date: str = None
    ) -> str:
        """
        Создать дневной отчет
        
        Returns:
            Путь к созданному файлу
        """
        if not date:
            date = datetime.now().strftime('%Y-%m-%d')
        
        # Получаем данные
        stats = database.get_today_stats()
        trades = database.get_recent_trades(limit=20)
        
        # Создаем документ
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('📊 ДНЕВНОЙ ТОРГОВЫЙ ОТЧЕТ', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата отчета
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'Дата: {date}')
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Пустая строка
        
        # ==== ОБЩАЯ СТАТИСТИКА ====
        doc.add_heading('📈 Общая статистика', level=1)
        
        # Создаем таблицу статистики
        stats_table = doc.add_table(rows=6, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        stats_data = [
            ('Всего сделок', stats.get('total_trades', 0)),
            ('Прибыльных', f"{stats.get('profitable_trades', 0)} ✅"),
            ('Убыточных', f"{stats.get('losing_trades', 0)} ❌"),
            ('Винрейт', f"{stats.get('win_rate', 0):.1f}%"),
            ('Общая прибыль', f"{stats.get('total_profit', 0):+.2f} USDT"),
            ('Объем торгов', f"{stats.get('total_volume', 0):.2f} USDT")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            # Жирный шрифт для значений
            cell = row.cells[1]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
        
        # ==== ДЕТАЛЬНАЯ АНАЛИТИКА ====
        doc.add_heading('🔍 Детальная аналитика', level=1)
        
        detail_para = doc.add_paragraph()
        detail_para.add_run(f"• Средняя прибыль: ").bold = True
        detail_para.add_run(f"{stats.get('avg_profit', 0):+.2f} USDT\n")
        
        detail_para.add_run(f"• Максимальная прибыль: ").bold = True
        detail_para.add_run(f"{stats.get('max_profit', 0):+.2f} USDT\n")
        
        detail_para.add_run(f"• Максимальный убыток: ").bold = True
        detail_para.add_run(f"{stats.get('min_profit', 0):+.2f} USDT\n")
        
        doc.add_paragraph()
        
        # ==== СПИСОК ПОСЛЕДНИХ СДЕЛОК ====
        if trades:
            doc.add_heading('📋 Последние сделки', level=1)
            
            trades_table = doc.add_table(rows=1, cols=7)
            trades_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = ['№', 'Сторона', 'Размер', 'Вход', 'Выход', 'Прибыль', 'Время']
            header_row = trades_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Данные сделок
            for idx, trade in enumerate(trades[:10], 1):
                row = trades_table.add_row()
                
                side_emoji = "🟢" if trade['side'] == 'long' else "🔴"
                profit_emoji = "💰" if trade['profit'] > 0 else "💸"
                
                row.cells[0].text = str(idx)
                row.cells[1].text = f"{side_emoji} {trade['side'].upper()}"
                row.cells[2].text = f"{trade['size']:.2f}"
                row.cells[3].text = f"{trade['entry_price']:.2f}"
                row.cells[4].text = f"{trade['exit_price']:.2f}"
                row.cells[5].text = f"{profit_emoji} {trade['profit']:+.2f}"
                
                close_time = trade['close_time'] if trade['close_time'] else '-'
                if close_time != '-':
                    close_time = datetime.strptime(close_time, '%Y-%m-%d %H:%M:%S').strftime('%H:%M')
                row.cells[6].text = close_time
        
        doc.add_paragraph()
        
        # ==== FOOTER ====
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f'\n\nОтчет сгенерирован автоматически\n'
            f'NADO DEX Trading Bot\n'
            f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Сохраняем документ
        filename = f"daily_report_{date}.docx"
        filepath = self.reports_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"✅ Дневной отчет создан: {filepath}")
        return str(filepath)
    
    def create_full_report(self, database) -> str:
        """
        Создать полный отчет за все время
        
        Returns:
            Путь к созданному файлу
        """
        # Получаем данные
        all_stats = database.get_all_time_stats()
        recent_trades = database.get_recent_trades(limit=50)
        daily_history = database.get_daily_stats_history(days=30)
        
        # Создаем документ
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('📊 ПОЛНЫЙ ТОРГОВЫЙ ОТЧЕТ', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run('Аналитика за весь период работы бота')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        
        # ==== ОБЩАЯ СТАТИСТИКА ЗА ВСЕ ВРЕМЯ ====
        doc.add_heading('📈 Общая статистика за все время', level=1)
        
        stats_table = doc.add_table(rows=9, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        stats_data = [
            ('Всего сделок', all_stats.get('total_trades', 0)),
            ('Прибыльных сделок', f"{all_stats.get('profitable_trades', 0)} ✅"),
            ('Убыточных сделок', f"{all_stats.get('losing_trades', 0)} ❌"),
            ('Винрейт', f"{all_stats.get('win_rate', 0):.2f}%"),
            ('Общая прибыль', f"{all_stats.get('total_profit', 0):+.2f} USDT"),
            ('Общий объем', f"{all_stats.get('total_volume', 0):,.2f} USDT"),
            ('Средняя прибыль', f"{all_stats.get('avg_profit', 0):+.2f} USDT"),
            ('Лучшая сделка', f"{all_stats.get('max_profit', 0):+.2f} USDT 🏆"),
            ('Худшая сделка', f"{all_stats.get('min_profit', 0):+.2f} USDT")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            cell = row.cells[1]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
        
        # ==== ИСТОРИЯ ПО ДНЯМ ====
        if daily_history:
            doc.add_heading('📅 История по дням (последние 30 дней)', level=1)
            
            history_table = doc.add_table(rows=1, cols=5)
            history_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = ['Дата', 'Сделки', 'Винрейт', 'Прибыль', 'Объем']
            header_row = history_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Данные по дням
            for day in daily_history[:15]:  # Последние 15 дней
                row = history_table.add_row()
                row.cells[0].text = day['date']
                row.cells[1].text = str(day['total_trades'])
                row.cells[2].text = f"{day['win_rate']:.1f}%"
                row.cells[3].text = f"{day['total_profit']:+.2f}"
                row.cells[4].text = f"{day['total_volume']:.2f}"
        
        doc.add_paragraph()
        
        # ==== ПОСЛЕДНИЕ СДЕЛКИ ====
        if recent_trades:
            doc.add_heading('📋 Последние 50 сделок', level=1)
            
            trades_table = doc.add_table(rows=1, cols=8)
            trades_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers = ['№', 'Дата', 'Сторона', 'Размер', 'Вход', 'Выход', 'Прибыль %', 'Прибыль']
            header_row = trades_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Данные сделок
            for idx, trade in enumerate(recent_trades, 1):
                row = trades_table.add_row()
                
                side_emoji = "🟢" if trade['side'] == 'long' else "🔴"
                
                row.cells[0].text = str(idx)
                
                # Дата
                close_time = trade['close_time'] if trade['close_time'] else '-'
                if close_time != '-':
                    close_time = datetime.strptime(close_time, '%Y-%m-%d %H:%M:%S').strftime('%d.%m %H:%M')
                row.cells[1].text = close_time
                
                row.cells[2].text = f"{side_emoji} {trade['side'].upper()}"
                row.cells[3].text = f"{trade['size']:.2f}"
                row.cells[4].text = f"{trade['entry_price']:.2f}"
                row.cells[5].text = f"{trade['exit_price']:.2f}" if trade['exit_price'] else '-'
                row.cells[6].text = f"{trade['profit_percent']:+.2f}%" if trade['profit_percent'] else '-'
                
                # Цвет прибыли
                profit_text = f"{trade['profit']:+.2f}"
                row.cells[7].text = profit_text
                
                if trade['profit'] > 0:
                    for paragraph in row.cells[7].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
                elif trade['profit'] < 0:
                    for paragraph in row.cells[7].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Красный
        
        doc.add_paragraph()
        
        # ==== FOOTER ====
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f'\n\nПолный отчет сгенерирован автоматически\n'
            f'NADO DEX Trading Bot\n'
            f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Сохраняем документ
        filename = f"full_report_{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
        filepath = self.reports_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"✅ Полный отчет создан: {filepath}")
        return str(filepath)
    
    def update_continuous_report(
        self,
        database,
        filename: str = "continuous_report.docx"
    ) -> str:
        """
        Обновить непрерывный отчет (добавить новые данные)
        Этот отчет постоянно обновляется с новой аналитикой
        
        Returns:
            Путь к обновленному файлу
        """
        filepath = self.reports_dir / filename
        
        # Если файл не существует - создаем новый
        if not filepath.exists():
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('📊 НЕПРЕРЫВНЫЙ ТОРГОВЫЙ ОТЧЕТ', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run('Этот документ постоянно обновляется')
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.italic = True
            
            doc.add_paragraph()
        else:
            # Открываем существующий документ
            doc = Document(str(filepath))
        
        # Добавляем разделитель
        doc.add_paragraph('_' * 80)
        
        # Добавляем обновление
        update_heading = doc.add_heading(
            f'📝 Обновление от {datetime.now().strftime("%Y-%m-%d %H:%M")}',
            level=2
        )
        
        # Текущая статистика
        stats = database.get_today_stats()
        all_stats = database.get_all_time_stats()
        
        # Статистика за сегодня
        doc.add_heading('Статистика за сегодня:', level=3)
        today_para = doc.add_paragraph()
        today_para.add_run(f"Сделок: {stats.get('total_trades', 0)} | ")
        today_para.add_run(f"Прибыль: {stats.get('total_profit', 0):+.2f} USDT | ")
        today_para.add_run(f"Винрейт: {stats.get('win_rate', 0):.1f}%")
        
        # Общая статистика
        doc.add_heading('Общая статистика:', level=3)
        total_para = doc.add_paragraph()
        total_para.add_run(f"Всего сделок: {all_stats.get('total_trades', 0)} | ")
        total_para.add_run(f"Общая прибыль: {all_stats.get('total_profit', 0):+.2f} USDT")
        
        doc.add_paragraph()
        
        # Сохраняем
        doc.save(str(filepath))
        logger.info(f"✅ Непрерывный отчет обновлен: {filepath}")
        return str(filepath)
